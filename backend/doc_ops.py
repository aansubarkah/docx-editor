import os, json, uuid, difflib
from typing import List, Tuple, Optional, Union
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_UNDERLINE
from models import Operation, OutlineItem
from utils import stable_paragraph_id, normalize_text

STORAGE_DIR = os.environ.get("STORAGE_DIR", os.path.join(os.path.dirname(__file__), "..", "storage"))
os.makedirs(STORAGE_DIR, exist_ok=True)

VERSIONS_DIR = os.path.join(STORAGE_DIR, "versions")
os.makedirs(VERSIONS_DIR, exist_ok=True)

def _file_path(file_id: str) -> str:
    return os.path.join(STORAGE_DIR, f"{file_id}.docx")

def _outline_path(file_id: str) -> str:
    return os.path.join(STORAGE_DIR, f"{file_id}.outline.json")

def save_new_doc(doc: Document) -> str:
    fid = str(uuid.uuid4())
    path = _file_path(fid)
    doc.save(path)
    # first outline
    outline = build_outline(fid)
    with open(_outline_path(fid), "w", encoding="utf-8") as f:
        json.dump([o.__dict__ for o in outline], f, ensure_ascii=False, indent=2)
    # version 1
    save_version(fid, path)
    return fid

def load_doc(file_id: str) -> Document:
    path = _file_path(file_id)
    if not os.path.exists(path):
        raise FileNotFoundError("file not found")
    return Document(path)

def create_document(title: str, body: Optional[str]) -> str:
    doc = Document()
    if title:
        doc.add_heading(title, 0)
    if body:
        for line in body.split("\n"):
            doc.add_paragraph(line)
    return save_new_doc(doc)

def _heading_level(p) -> int:
    try:
        name = (p.style.name or "").lower()
    except Exception:
        name = ""
    if name.startswith("heading"):
        try:
            parts = name.split()
            lvl = int(parts[1]) if len(parts)>1 else 1
            return max(1, min(6, lvl))
        except Exception:
            return 1
    return 0

def build_outline(file_id: str) -> List[OutlineItem]:
    doc = load_doc(file_id)
    outline: List[OutlineItem] = []
    for i, p in enumerate(doc.paragraphs):
        lvl = _heading_level(p)
        pid = stable_paragraph_id(p.text, i, lvl)
        outline.append(OutlineItem(paragraph_id=pid, text=p.text or "", level=lvl))
    return outline

def load_outline(file_id: str) -> List[OutlineItem]:
    path = _outline_path(file_id)
    if os.path.exists(path):
        try:
            data = json.load(open(path, "r", encoding="utf-8"))
            return [OutlineItem(**x) for x in data]
        except Exception:
            pass
    return build_outline(file_id)

def _find_insert_index_by_anchor(doc: Document, anchor_pid: str) -> Optional[int]:
    # Recompute live paragraph IDs and match
    current = []
    for i, p in enumerate(doc.paragraphs):
        lvl = _heading_level(p)
        pid = stable_paragraph_id(p.text, i, lvl)
        current.append((i, pid))
    for idx, pid in current:
        if pid == anchor_pid:
            return idx + 1
    return None

def _get_default_font(doc: Document) -> Optional[str]:
    """Extract the most common font from existing paragraphs"""
    font_counts = {}
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name:
                font_counts[run.font.name] = font_counts.get(run.font.name, 0) + 1
    if font_counts:
        return max(font_counts, key=font_counts.get)
    return None

def _copy_paragraph_formatting(source_para, target_para):
    """Copy formatting from source paragraph to target paragraph"""
    if not source_para.runs:
        return

    # Get font from first run of source
    source_run = source_para.runs[0]
    target_text = target_para.text

    # Clear target and recreate with formatting
    target_para.clear()
    new_run = target_para.add_run(target_text)

    # Copy font properties
    if source_run.font.name:
        new_run.font.name = source_run.font.name
    if source_run.font.size:
        new_run.font.size = source_run.font.size
    new_run.font.bold = source_run.font.bold
    new_run.font.italic = source_run.font.italic

def apply_operations(file_id: str, operations: List[Operation]) -> Tuple[str, List[OutlineItem]]:
    doc = load_doc(file_id)

    # Get default font from document
    default_font = _get_default_font(doc)
    # Find a reference paragraph for formatting
    ref_paragraph = None
    for p in doc.paragraphs:
        if p.runs and _heading_level(p) == 0:  # Get a normal paragraph
            ref_paragraph = p
            break

    for op in operations:
        if op.type == "add_heading":
            level = 1 if op.level is None else max(1, min(6, int(op.level)))
            text = op.text or ""
            if op.after_paragraph_id:
                idx = _find_insert_index_by_anchor(doc, op.after_paragraph_id)
                if idx is None: 
                    p = doc.add_paragraph()
                    p.style = f"Heading {level}"
                    p.text = text
                else:
                    # python-docx lacks "insert at index"; rebuild lightly
                    texts = [(p.text, _heading_level(p)) for p in doc.paragraphs]
                    texts.insert(idx, (text, level))
                    new_doc = Document()
                    for t, lvl in texts:
                        if lvl>0:
                            h = new_doc.add_heading(t, level=lvl)
                        else:
                            new_doc.add_paragraph(t)
                    doc = new_doc
            else:
                doc.add_heading(text, level=level)

        elif op.type == "add_paragraph":
            text = op.text or ""
            if op.after_paragraph_id:
                idx = _find_insert_index_by_anchor(doc, op.after_paragraph_id)
                if idx is None:
                    new_para = doc.add_paragraph(text)
                    if ref_paragraph:
                        _copy_paragraph_formatting(ref_paragraph, new_para)
                else:
                    texts = [(p.text, _heading_level(p)) for p in doc.paragraphs]
                    texts.insert(idx, (text, 0))
                    new_doc = Document()
                    for t, lvl in texts:
                        if lvl>0:
                            new_doc.add_heading(t, level=lvl)
                        else:
                            new_para = new_doc.add_paragraph(t)
                            if ref_paragraph:
                                _copy_paragraph_formatting(ref_paragraph, new_para)
                    doc = new_doc
            else:
                new_para = doc.add_paragraph(text)
                if ref_paragraph:
                    _copy_paragraph_formatting(ref_paragraph, new_para)

        elif op.type == "replace_text":
            if not op.find:
                continue
            for p in doc.paragraphs:
                if p.text and op.find in p.text:
                    # Preserve formatting by replacing text in runs
                    new_text = p.text.replace(op.find, op.replace or "")

                    # Store original formatting
                    original_runs = []
                    for run in p.runs:
                        original_runs.append({
                            'font_name': run.font.name,
                            'font_size': run.font.size,
                            'bold': run.font.bold,
                            'italic': run.font.italic
                        })

                    # Replace text
                    p.clear()
                    if original_runs:
                        run = p.add_run(new_text)
                        # Apply first run's formatting
                        fmt = original_runs[0]
                        if fmt['font_name']:
                            run.font.name = fmt['font_name']
                        if fmt['font_size']:
                            run.font.size = fmt['font_size']
                        run.font.bold = fmt['bold']
                        run.font.italic = fmt['italic']
                    else:
                        p.add_run(new_text)

        elif op.type == "insert_table":
            rows = op.rows or (len(op.data) if op.data else 2)
            cols = op.cols or (len(op.data[0]) if (op.data and len(op.data)>0) else 2)

            # Insert table with styling
            from docx.shared import Pt
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls

            table = doc.add_table(rows=rows, cols=cols)

            # Try to apply table style if available
            try:
                table.style = 'Table Grid'
            except KeyError:
                # If Table Grid doesn't exist, add borders manually
                tbl = table._element
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = parse_xml(r'<w:tblPr ' + nsdecls('w') + '/>')
                    tbl.insert(0, tblPr)
                # Add borders
                tblBorders = parse_xml(r'<w:tblBorders ' + nsdecls('w') + '>'
                    r'<w:top w:val="single" w:sz="4"/>'
                    r'<w:left w:val="single" w:sz="4"/>'
                    r'<w:bottom w:val="single" w:sz="4"/>'
                    r'<w:right w:val="single" w:sz="4"/>'
                    r'<w:insideH w:val="single" w:sz="4"/>'
                    r'<w:insideV w:val="single" w:sz="4"/>'
                    r'</w:tblBorders>')
                tblPr.append(tblBorders)

            # Populate table data
            if op.data:
                for i in range(min(rows, len(op.data))):
                    for j in range(min(cols, len(op.data[i]))):
                        cell = table.cell(i, j)
                        cell.text = str(op.data[i][j])

                        # Apply font formatting to match document
                        if default_font and cell.paragraphs:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = default_font
                                    run.font.size = Pt(11)

            # Add header row styling if requested
            if op.add_header_row and rows > 0:
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

        elif op.type == "edit_table":
            # Edit existing table
            if op.table_index is not None and op.table_index < len(doc.tables):
                table = doc.tables[op.table_index]

                # Edit specific cell
                if op.cell_row is not None and op.cell_col is not None and op.cell_text is not None:
                    if op.cell_row < len(table.rows) and op.cell_col < len(table.columns):
                        cell = table.cell(op.cell_row, op.cell_col)

                        # Preserve formatting
                        original_font = None
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            original_font = {
                                'name': cell.paragraphs[0].runs[0].font.name,
                                'size': cell.paragraphs[0].runs[0].font.size,
                                'bold': cell.paragraphs[0].runs[0].font.bold,
                                'italic': cell.paragraphs[0].runs[0].font.italic
                            }

                        cell.text = op.cell_text

                        # Reapply formatting
                        if original_font and cell.paragraphs:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if original_font['name']:
                                        run.font.name = original_font['name']
                                    if original_font['size']:
                                        run.font.size = original_font['size']
                                    run.font.bold = original_font['bold']
                                    run.font.italic = original_font['italic']

                # Populate all cells if data provided
                elif op.data:
                    for i in range(min(len(table.rows), len(op.data))):
                        for j in range(min(len(table.columns), len(op.data[i]))):
                            cell = table.cell(i, j)
                            cell.text = str(op.data[i][j])

        elif op.type == "remove_table":
            # Remove table by index
            if op.table_index is not None and op.table_index < len(doc.tables):
                table = doc.tables[op.table_index]
                # Get the table element and remove it from the document
                tbl_element = table._element
                tbl_element.getparent().remove(tbl_element)

        elif op.type == "remove_paragraph":
            # Remove paragraph by paragraph_id or by text match
            if op.after_paragraph_id:
                # Remove by stable paragraph ID
                paragraphs_to_remove = []
                for i, p in enumerate(doc.paragraphs):
                    lvl = _heading_level(p)
                    pid = stable_paragraph_id(p.text, i, lvl)
                    if pid == op.after_paragraph_id:
                        paragraphs_to_remove.append(p)

                for p in paragraphs_to_remove:
                    p_element = p._element
                    p_element.getparent().remove(p_element)

            elif op.find:
                # Remove by text match (all paragraphs containing the text)
                paragraphs_to_remove = []
                for p in doc.paragraphs:
                    if p.text and op.find in p.text:
                        paragraphs_to_remove.append(p)

                for p in paragraphs_to_remove:
                    p_element = p._element
                    p_element.getparent().remove(p_element)

    # Save as new version (incremental)
    new_id = file_id  # keep same id; version separately
    path = _file_path(new_id)
    doc.save(path)

    # Update outline and persist
    outline = build_outline(new_id)
    with open(_outline_path(new_id), "w", encoding="utf-8") as f:
        json.dump([o.__dict__ for o in outline], f, ensure_ascii=False, indent=2)

    # Add version snapshot
    save_version(new_id, path)

    return new_id, outline

def save_version(file_id: str, src_path: str) -> str:
    versions_dir = os.path.join(VERSIONS_DIR, file_id)
    os.makedirs(versions_dir, exist_ok=True)
    n = len([f for f in os.listdir(versions_dir) if f.endswith(".docx")])
    vname = f"v{n+1}.docx"
    dst = os.path.join(versions_dir, vname)
    import shutil
    shutil.copy2(src_path, dst)
    return vname

def list_versions(file_id: str):
    versions_dir = os.path.join(VERSIONS_DIR, file_id)
    if not os.path.exists(versions_dir): return []
    items = sorted([f for f in os.listdir(versions_dir) if f.endswith(".docx")])
    return items

# --------- Redline-style compare (visual diff) ---------
def _paragraph_texts(doc: Document):
    return [p.text or "" for p in doc.paragraphs]

def _apply_run_style(run, color: Optional[Tuple[int,int,int]]=None, underline=False, strike=False):
    if color:
        run.font.color.rgb = RGBColor(*color)
    if underline:
        run.underline = True
    if strike:
        run.font.strike = True

def _compose_diff_doc(base_texts: List[str], revised_texts: List[str]) -> Document:
    # Very simple: line-based compare with word-diff inside changed lines
    out = Document()
    out.add_heading("Redline (visual) compare", level=1)
    sm = difflib.SequenceMatcher(a=base_texts, b=revised_texts)
    for opcode, i1, i2, j1, j2 in sm.get_opcodes():
        if opcode == "equal":
            for k in range(i1, i2):
                out.add_paragraph(base_texts[k])
        elif opcode == "delete":
            for k in range(i1, i2):
                p = out.add_paragraph()
                run = p.add_run(base_texts[k])
                _apply_run_style(run, color=(220,0,0), strike=True)
        elif opcode == "insert":
            for k in range(j1, j2):
                p = out.add_paragraph()
                run = p.add_run(revised_texts[k])
                _apply_run_style(run, color=(0,140,0), underline=True)
        elif opcode == "replace":
            # word-level diff
            for a_line, b_line in zip(base_texts[i1:i2], revised_texts[j1:j2] or [""]*(i2-i1)):
                p = out.add_paragraph()
                aw = a_line.split()
                bw = b_line.split()
                wsm = difflib.SequenceMatcher(a=aw, b=bw)
                for op, ai1, ai2, bj1, bj2 in wsm.get_opcodes():
                    if op == "equal":
                        p.add_run(" " + " ".join(aw[ai1:ai2]))
                    elif op == "delete":
                        r = p.add_run(" " + " ".join(aw[ai1:ai2]))
                        _apply_run_style(r, color=(220,0,0), strike=True)
                    elif op == "insert":
                        r = p.add_run(" " + " ".join(bw[bj1:bj2]))
                        _apply_run_style(r, color=(0,140,0), underline=True)
                    elif op == "replace":
                        r1 = p.add_run(" " + " ".join(aw[ai1:ai2]))
                        _apply_run_style(r1, color=(220,0,0), strike=True)
                        r2 = p.add_run(" " + " ".join(bw[bj1:bj2]))
                        _apply_run_style(r2, color=(0,140,0), underline=True)
    return out

def redline_compare(base_id: str, revised_id: str) -> str:
    base = load_doc(base_id)
    rev = load_doc(revised_id)
    base_texts = _paragraph_texts(base)
    revised_texts = _paragraph_texts(rev)
    out = _compose_diff_doc(base_texts, revised_texts)
    out_id = f"compare-{uuid.uuid4()}"
    out.save(os.path.join(STORAGE_DIR, f"{out_id}.docx"))
    return out_id
